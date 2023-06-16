from dotenv import load_dotenv
import os
from PyPDF2 import PdfReader
import docx
from docx import Document
from docx.shared import Inches
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from langchain.callbacks import get_openai_callback
from langchain.chat_models import ChatOpenAI

load_dotenv()

# Create a new Word document
document = Document()

#### Define Section Prompts ####
admin_items = "Using the context documents, write a response to these items and provide quotes where possible.: " \
                        "Duration: <How long is the grant for?> " \
                        "Closing date: <What is the closing date?> " \
                        "Enquiries: <What is the phone/email address for any enquiries?>" \
                        "Letter of Support: <What if any support documents are required?>"

project_alignment = "Using the context documents, write a response to these items and provide quotes where possible.: " \
                        "Scheme Goals: <What are the major goals of this scheme?> " \
                        "Target Audience: <Who is the target audience for this scheme?> " \
                        "Funding body Mission: <What is the overall mission of the funding body?>"

budget = "Using the context documents, write a response to these items and provide quotes where possible.: " \
            "Ineligible project costs: <What items are not covered by this scheme> " \
            "Eligible project costs: <What are items that can be requested in this scheme?> " \
            "Maximum Amount: <What is the maximum amount that will be awarded?>"

eligibility_items = "Using the context documents, write a response to these items and provide quotes where possible.: " \
                        "PhD Requirement: <Is a PhD required for this application and is there a cut-off date for when they were awarded their PhD?> " \
                        "Citizenship Requirement: <Is there a requirement to be an Australian Citizen or Permanent resident?> " \
                        "Employment Requirement: <Is there a requirement to be currently employed at time of grant application?>"

formatting = "Using the context documents, write a response to these items and provide quotes where possible.:" \
                "Page limit: <What is/Is there a page limit?>" \
                "Formatting Requirements: <What is/Is there formatting requirements?>" \
                "Formatting Other: <What other important information is required related to the format?>"

other_important_information = "In addition to project alignment, formatting, budget and eligibility items, please write three points of information that would be important to applicants. " \
                              "Separate each response with a new line"

############ TEXT LOADERS ############
# Functions to read different file types
def read_pdf(file_path):
    with open(file_path, "rb") as file:
        pdf_reader = PdfReader(file)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
    return text

def read_word(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_txt(file_path):
    with open(file_path, "r") as file:
        text = file.read()
    return text

def read_documents_from_directory(directory):
    combined_text = ""
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename.endswith(".pdf"):
            combined_text += read_pdf(file_path)
        elif filename.endswith(".docx"):
            combined_text += read_word(file_path)
        elif filename.endswith(".txt"):
            combined_text += read_txt(file_path)
    return combined_text

########## DOCUMENT QUERY ##############
def query_docs(question):
    train_directory = 'train_files/'
    text = read_documents_from_directory(train_directory)

    # split into chunks
    char_text_splitter = CharacterTextSplitter(separator="\n", chunk_size=4000,
                                               chunk_overlap=400, length_function=len)

    text_chunks = char_text_splitter.split_text(text)

    # create embeddings
    embeddings = OpenAIEmbeddings()
    docsearch = FAISS.from_texts(text_chunks, embeddings)

    llm = ChatOpenAI(model_name="gpt-3.5-turbo-16k")
    chain = load_qa_chain(llm, chain_type="stuff")

    docs = docsearch.similarity_search(question)
    response = chain.run(input_documents=docs, question=question)

    print(response)
    output_response = response.strip().split("\n")
    with get_openai_callback() as cb:
        response = chain.run(input_documents=docs, question=question)
        print(cb)
    return output_response

############ ADD SECTION BASED ON RESPONSES ###############

def add_section_to_doc(response_list, section_name):
    # Add section to the document using the responses imported in list format.
    document.add_heading(section_name, level=2)

    # Iterate over the response list
    for response in response_list:
        try:
            # Add a paragraph for each response
            document.add_paragraph(str(response))
        except Exception as e:
            # Handle the exception
            print(f"An error occurred: {e}")
            # Add an escape mechanism or perform any necessary actions
            break

def make_document(scheme):

    # Access the first section of the document
    section = document.sections[0]

    # Set page size to A4 (210mm x 297mm)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Add a header
    header = section.header

    # Create a paragraph in the header and set its alignment
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add content to the header paragraph
    run = header_paragraph.add_run()
    run.add_text('Auto-Generated Review Document')

    # Add a footer
    footer = section.footer

    # Create a paragraph in the footer and set its alignment
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add content to the footer paragraph
    run = footer_paragraph.add_run()
    run.add_text('Office for Research - Griffith University')

    # Add Headers
    document.add_heading('Auto-generated Review Document for ' + scheme, level=1)

    add_section_to_doc(query_docs(admin_items), "Admin Items")
    add_section_to_doc(query_docs(project_alignment), "Project Alignment")
    add_section_to_doc(query_docs(eligibility_items), "Eligibility Items")
    add_section_to_doc(query_docs(budget), "Budget")
    add_section_to_doc(query_docs(formatting), "Formatting")
    add_section_to_doc(query_docs(other_important_information), "Other Important Things to Note")

    document.save(scheme + ".docx")

make_document("Scheme Name")
